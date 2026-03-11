import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile
import os
import fitz  # PyMuPDF
import io
import re


def extract_text_from_pdf(pdf_bytes: bytes) -> str:
    """Extract plain text from PDF bytes."""
    with fitz.open(stream=pdf_bytes, filetype="pdf") as doc:
        pages = [page.get_text() for page in doc]
    return "\n\n".join(pages)


def extract_text_from_docx(docx_bytes: bytes) -> list[dict]:
    """
    Extract structured text from DOCX bytes.
    Returns a list of segments with type and text.
    """
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp.write(docx_bytes)
        tmp_path = tmp.name

    try:
        doc = docx.Document(tmp_path)
        segments = []
        segment_id = 1
        list_counters = {}
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            style_name = para.style.name.lower() if para.style else ""
            
            if style_name.startswith('heading 1'):
                seg_type = "h1"
            elif style_name.startswith('heading 2'):
                seg_type = "h2"
            elif style_name.startswith('heading 3'):
                seg_type = "h3"
            elif 'list' in style_name or 'bullet' in style_name:
                seg_type = "li"
            elif 'number' in style_name or 'ordered' in style_name:
                num_id = para._element.pPr.numPr.numId.val if para._element.pPr is not None and para._element.pPr.numPr is not None else None
                if num_id is not None:
                    if num_id not in list_counters:
                        list_counters[num_id] = 1
                    else:
                        list_counters[num_id] += 1
                    seg_type = "ol"
                else:
                    seg_type = "li"
            elif para.runs and any(r.bold for r in para.runs if r.text.strip()):
                all_bold = all(r.bold for r in para.runs if r.text.strip())
                if all_bold and len(text) < 100:
                    seg_type = "h3"
                else:
                    seg_type = "li"
            elif para.paragraph_format.left_indent is not None and para.paragraph_format.left_indent > 0:
                seg_type = "li"
            else:
                seg_type = "p"
            
            extra_data = {}
            if seg_type == "ol" and num_id is not None:
                extra_data["number"] = list_counters[num_id]
            
            segments.append({
                "id": segment_id,
                "type": seg_type,
                "text": text,
                **extra_data
            })
            segment_id += 1
        
        return segments
    finally:
        os.unlink(tmp_path)


def extract_text_plain_from_docx(docx_bytes: bytes) -> str:
    """Extract plain text from Word document bytes."""
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp.write(docx_bytes)
        tmp_path = tmp.name

    try:
        doc = docx.Document(tmp_path)
        paragraphs = [para.text for para in doc.paragraphs]
        return "\n\n".join(paragraphs)
    finally:
        os.unlink(tmp_path)


def create_translated_docx(translated_text: str, source_filename: str = "translated") -> bytes:
    """Generate a Word document from translated text. Returns DOCX as bytes."""
    doc = docx.Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    for line in translated_text.split('\n'):
        doc.add_paragraph(line)
    
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        doc.save(tmp.name)
        tmp_path = tmp.name

    try:
        with open(tmp_path, "rb") as f:
            docx_bytes = f.read()
    finally:
        os.unlink(tmp_path)
    
    return docx_bytes


def create_frozen_template_pdf(
    original_file: bytes,
    segments: list[dict],
    target_lang: str,
    doc_title: str = "Translated Document",
    doc_ref: str = "DOC-001"
) -> bytes:
    """
    Create a PDF with proper HTML-like formatting using PyMuPDF.
    Preserves document structure with styled headings, sections, and bullet points.
    """
    # Page dimensions (A4)
    page_width = 595
    page_height = 842
    margin_left = 50
    margin_right = 50
    margin_top = 60
    margin_bottom = 50
    content_width = page_width - margin_left - margin_right
    
    # Create PDF document
    doc = fitz.open()
    page = doc.new_page(width=page_width, height=page_height)
    
    # Colors
    color_black = (0, 0, 0)
    color_dark_blue = (0.1, 0.2, 0.4)
    color_gray = (0.5, 0.5, 0.5)
    
    y_position = margin_top
    
    # Helper function to wrap text
    def wrap_text(text, max_width, fontsize):
        words = text.split()
        lines = []
        current_line = []
        
        for word in words:
            test_line = ' '.join(current_line + [word])
            estimated_width = len(test_line) * fontsize * 0.4
            
            if estimated_width <= max_width:
                current_line.append(word)
            else:
                if current_line:
                    lines.append(' '.join(current_line))
                current_line = [word]
        
        if current_line:
            lines.append(' '.join(current_line))
        
        return lines if lines else [""]
    
    # Helper function to check if new page needed
    def check_new_page(required_height):
        nonlocal y_position, page
        if y_position + required_height > page_height - margin_bottom:
            page = doc.new_page(width=page_width, height=page_height)
            y_position = margin_top
            return True
        return False
    
    # Helper function to insert centered text
    def insert_centered_text(page, y, text, fontsize, color):
        text_width = len(text) * fontsize * 0.4
        x = (page_width - text_width) / 2
        page.insert_text(fitz.Point(x, y), text, fontsize=fontsize, color=color)
    
    # ====== HEADER SECTION ======
    # Title
    check_new_page(30)
    insert_centered_text(page, y_position, doc_title, 20, color_dark_blue)
    y_position += 25
    
    # Subtitle
    subtitle = f"Reference: {doc_ref} | Language: {target_lang}"
    insert_centered_text(page, y_position, subtitle, 10, color_gray)
    y_position += 20
    
    # Horizontal line
    page.draw_line(fitz.Point(margin_left, y_position), fitz.Point(page_width - margin_right, y_position), color=(0.7, 0.7, 0.7), width=0.5)
    y_position += 15
    
    # ====== CONTENT SECTION ======
    for seg in segments:
        text = seg.get("translated_text", seg.get("text", ""))
        if not text:
            continue
        
        seg_type = seg.get("type", "p")
        
        # Apply formatting based on segment type
        if seg_type == "h1":
            # Heading 1: Large, bold, dark blue
            check_new_page(25)
            page.insert_text(fitz.Point(margin_left, y_position), text, fontsize=18, color=color_dark_blue)
            y_position += 30
            
        elif seg_type == "h2":
            # Heading 2: Medium-large, bold
            check_new_page(22)
            page.insert_text(fitz.Point(margin_left, y_position), text, fontsize=15, color=color_black)
            y_position += 24
            
        elif seg_type == "h3":
            # Heading 3: Medium, bold
            check_new_page(20)
            page.insert_text(fitz.Point(margin_left, y_position), text, fontsize=13, color=color_black)
            y_position += 20
            
        elif seg_type == "li":
            # List item: bullet with indentation
            indent = 20
            wrapped = wrap_text(text, content_width - indent, 11)
            
            for i, line_text in enumerate(wrapped):
                check_new_page(16)
                bullet_x = margin_left + 5
                text_x = margin_left + indent
                
                if i == 0:
                    page.insert_text(fitz.Point(bullet_x, y_position), "•", fontsize=12, color=color_black)
                
                page.insert_text(fitz.Point(text_x, y_position), line_text, fontsize=11, color=color_black)
                y_position += 14
            y_position += 4
            
        elif seg_type == "ol":
            # Ordered list: number with indentation
            indent = 25
            number = seg.get("number", 1)
            wrapped = wrap_text(text, content_width - indent, 11)
            
            for i, line_text in enumerate(wrapped):
                check_new_page(16)
                num_x = margin_left + 10
                text_x = margin_left + indent
                
                if i == 0:
                    page.insert_text(fitz.Point(num_x, y_position), f"{number}.", fontsize=11, color=color_black)
                
                page.insert_text(fitz.Point(text_x, y_position), line_text, fontsize=11, color=color_black)
                y_position += 14
            y_position += 4
            
        else:
            # Regular paragraph
            wrapped = wrap_text(text, content_width, 11)
            
            for line_text in wrapped:
                check_new_page(16)
                page.insert_text(fitz.Point(margin_left, y_position), line_text, fontsize=11, color=color_black)
                y_position += 14
            y_position += 8
    
    # Add page numbers
    for p in doc:
        footer_text = f"Page {p.number}"
        text_width = len(footer_text) * 9 * 0.4
        x = (page_width - text_width) / 2
        page_rect = p.rect
        footer_y = page_rect.height - 30
        p.insert_text(fitz.Point(x, footer_y), footer_text, fontsize=9, color=color_gray)
    
    return doc.tobytes()


def create_translated_pdf(translated_text: str, source_filename: str = "translated") -> bytes:
    """
    Generate a PDF from translated text with proper formatting using PyMuPDF.
    Automatically detects headings, subheadings, and list items.
    """
    # Page dimensions (A4)
    page_width = 595
    page_height = 842
    margin_left = 50
    margin_right = 50
    margin_top = 60
    margin_bottom = 50
    content_width = page_width - margin_left - margin_right
    
    # Create PDF document
    doc = fitz.open()
    page = doc.new_page(width=page_width, height=page_height)
    
    # Colors
    color_black = (0, 0, 0)
    color_dark_blue = (0.1, 0.2, 0.4)
    color_gray = (0.5, 0.5, 0.5)
    
    y_position = margin_top
    
    # Helper function to wrap text
    def wrap_text(text, max_width, fontsize):
        words = text.split()
        lines = []
        current_line = []
        
        for word in words:
            test_line = ' '.join(current_line + [word])
            estimated_width = len(test_line) * fontsize * 0.4
            
            if estimated_width <= max_width:
                current_line.append(word)
            else:
                if current_line:
                    lines.append(' '.join(current_line))
                current_line = [word]
        
        if current_line:
            lines.append(' '.join(current_line))
        
        return lines if lines else [""]
    
    # Helper function to check if new page needed
    def check_new_page(required_height):
        nonlocal y_position, page
        if y_position + required_height > page_height - margin_bottom:
            page = doc.new_page(width=page_width, height=page_height)
            y_position = margin_top
            return True
        return False
    
    # Helper function to insert centered text
    def insert_centered_text(page, y, text, fontsize, color):
        text_width = len(text) * fontsize * 0.4
        x = (page_width - text_width) / 2
        page.insert_text(fitz.Point(x, y), text, fontsize=fontsize, color=color)
    
    # Add header
    check_new_page(30)
    title = source_filename.replace(".pdf", "").replace("_", " ").title()
    insert_centered_text(page, y_position, title, 20, color_dark_blue)
    y_position += 25
    
    # Horizontal line
    page.draw_line(fitz.Point(margin_left, y_position), fitz.Point(page_width - margin_right, y_position), color=(0.7, 0.7, 0.7), width=0.5)
    y_position += 15
    
    # Parse and format the text
    lines = translated_text.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            y_position += 8
            continue
        
        # Detect formatting based on line content
        is_heading = False
        heading_level = 0
        
        # Check for markdown-style headings
        if line.startswith('# '):
            heading_level = 1
            line = line[2:].strip()
            is_heading = True
        elif line.startswith('## '):
            heading_level = 2
            line = line[3:].strip()
            is_heading = True
        elif line.startswith('### '):
            heading_level = 3
            line = line[4:].strip()
            is_heading = True
        # Check for ALL CAPS lines (often headings)
        elif line.isupper() and len(line) < 80 and not any(c.isdigit() for c in line):
            heading_level = 2
            is_heading = True
        # Check for lines ending with colon (subheadings)
        elif line.endswith(':') and len(line) < 60:
            heading_level = 3
            is_heading = True
        
        # Check for bullet points
        is_bullet = line.startswith('• ') or line.startswith('- ') or line.startswith('* ')
        if is_bullet:
            line = line[2:].strip()
        
        # Check for numbered lists
        is_numbered = bool(re.match(r'^\d+[\.)]\s', line))
        if is_numbered:
            line = re.sub(r'^\d+[\.)]\s', '', line).strip()
        
        # Format based on type
        if is_heading:
            check_new_page(25)
            if heading_level == 1:
                page.insert_text(fitz.Point(margin_left, y_position), line, fontsize=18, color=color_dark_blue)
                y_position += 30
            elif heading_level == 2:
                page.insert_text(fitz.Point(margin_left, y_position), line, fontsize=15, color=color_black)
                y_position += 24
            else:
                page.insert_text(fitz.Point(margin_left, y_position), line, fontsize=13, color=color_black)
                y_position += 20
                
        elif is_bullet:
            indent = 20
            wrapped = wrap_text(line, content_width - indent, 11)
            
            for i, line_text in enumerate(wrapped):
                check_new_page(16)
                bullet_x = margin_left + 5
                text_x = margin_left + indent
                
                if i == 0:
                    page.insert_text(fitz.Point(bullet_x, y_position), "•", fontsize=12, color=color_black)
                
                page.insert_text(fitz.Point(text_x, y_position), line_text, fontsize=11, color=color_black)
                y_position += 14
            y_position += 4
            
        elif is_numbered:
            indent = 25
            wrapped = wrap_text(line, content_width - indent, 11)
            
            for i, line_text in enumerate(wrapped):
                check_new_page(16)
                num_x = margin_left + 10
                text_x = margin_left + indent
                
                if i == 0:
                    page.insert_text(fitz.Point(num_x, y_position), "1.", fontsize=11, color=color_black)
                
                page.insert_text(fitz.Point(text_x, y_position), line_text, fontsize=11, color=color_black)
                y_position += 14
            y_position += 4
            
        else:
            # Regular paragraph
            wrapped = wrap_text(line, content_width, 11)
            
            for line_text in wrapped:
                check_new_page(16)
                page.insert_text(fitz.Point(margin_left, y_position), line_text, fontsize=11, color=color_black)
                y_position += 14
            y_position += 8
    
    # Add page numbers
    for p in doc:
        footer_text = f"Page {p.number}"
        text_width = len(footer_text) * 9 * 0.4
        x = (page_width - text_width) / 2
        page_rect = p.rect
        footer_y = page_rect.height - 30
        p.insert_text(fitz.Point(x, footer_y), footer_text, fontsize=9, color=color_gray)
    
    return doc.tobytes()
